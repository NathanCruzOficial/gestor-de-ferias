from docx import Document
import os

def preencher_docx(dados, tabela_dados):
    modelo_path = os.path.abspath("app/models/docs/doc-fry-model.docx")
    output_dir = os.path.abspath("documentos_gerados")
    output_path = os.path.join(output_dir, "relatorio_preenchido.docx")

    # Criar diretório se não existir
    os.makedirs(output_dir, exist_ok=True)

    # Verificar se o arquivo modelo existe
    if not os.path.exists(modelo_path):
        raise FileNotFoundError(f"Arquivo modelo não encontrado: {modelo_path}")

    # Carregar o documento modelo
    doc = Document(modelo_path)

    # Substituir placeholders nos parágrafos
    for paragrafo in doc.paragraphs:
        for chave, valor in dados.items():
            if f"{{{{{chave}}}}}" in paragrafo.text:
                paragrafo.text = paragrafo.text.replace(f"{{{{{chave}}}}}", valor)

    # Substituir placeholders no rodapé
    for section in doc.sections:
        footer = section.footer
        for paragrafo in footer.paragraphs:
            for chave, valor in dados.items():
                if f"{{{{{chave}}}}}" in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace(f"{{{{{chave}}}}}", valor)

    # Substituir placeholders no rodapé
    for section in doc.sections:
        header = section.header
        for paragrafo in header.paragraphs:
            for chave, valor in dados.items():
                if f"{{{{{chave}}}}}" in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace(f"{{{{{chave}}}}}", valor)

      # Preencher a tabela com os dados
    for tabela in doc.tables:
        if len(tabela.rows) > 0:  # Garante que a tabela tenha pelo menos um cabeçalho
            cabecalho = tabela.rows[0].cells  # Assume que a primeira linha é o cabeçalho
            
            # Verificar o número de colunas na tabela
            num_colunas = len(cabecalho)
            
            # Adiciona cada linha de dados verificando o número de colunas
            for item in tabela_dados:
                linha = tabela.add_row().cells
                valores = [
                    item.get("state", ""), 
                    item.get("om", ""), 
                    item.get("p/g", ""),  
                    item.get("nome", ""),
                    item.get("data_inicio", ""), 
                    str(item.get("dias", "")),
                    item.get("data_retorno", ""),
                    item.get("contato", "")]

                for i in range(min(len(valores), num_colunas)):  # Garante que não acesse colunas inexistentes
                    linha[i].text = valores[i]

    # Salvar o arquivo
    doc.save(output_path)

    # Verificar se o arquivo foi criado corretamente
    if not os.path.exists(output_path):
        raise FileNotFoundError(f"Erro ao criar o arquivo: {output_path}")

    # Retornar o arquivo como download
    return output_path